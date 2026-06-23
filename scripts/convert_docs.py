import pandas as pd
import docx
import os

files = [
    r"d:\Minimax\docs\Bidder_Copy_Clarification_Qs_Healthcare_AI_SBS10525_BATCH 2_03_06_26.xlsx",
    r"d:\Minimax\docs\Healthcare AI Solutions Bidder Declarations Declaration of Interest.docx",
    r"d:\Minimax\docs\Healthcare AI Solutions Bidders Instructions and Guidance V2.docx",
    r"d:\Minimax\docs\Healthcare AI Solutions Commercial Schedule - Lot 7.xlsx"
]

output_dir = r"d:\Minimax\nhs_sbs_cq_portal\data\markdown_docs"
os.makedirs(output_dir, exist_ok=True)

for file_path in files:
    filename = os.path.basename(file_path)
    base_name, ext = os.path.splitext(filename)
    out_file = os.path.join(output_dir, base_name + ".md")
    
    print(f"Processing {filename}...")
    try:
        if ext.lower() == '.xlsx':
            xls = pd.ExcelFile(file_path)
            with open(out_file, 'w', encoding='utf-8') as f:
                f.write(f"# {base_name}\n\n")
                for sheet in xls.sheet_names:
                    f.write(f"## Sheet: {sheet}\n\n")
                    df = pd.read_excel(file_path, sheet_name=sheet)
                    f.write(df.to_markdown(index=False))
                    f.write("\n\n")
            print(f"-> Successfully converted {filename} to MD.")
            
        elif ext.lower() == '.docx':
            doc = docx.Document(file_path)
            with open(out_file, 'w', encoding='utf-8') as f:
                f.write(f"# {base_name}\n\n")
                
                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        f.write(para.text + "\n\n")
                
                # Extract tables
                for i, table in enumerate(doc.tables):
                    f.write(f"\n### Table {i+1}\n\n")
                    if len(table.rows) > 0:
                        # Header
                        f.write("| " + " | ".join([cell.text.replace("\n", " ") for cell in table.rows[0].cells]) + " |\n")
                        f.write("|" + "|".join(["---" for _ in table.rows[0].cells]) + "|\n")
                        # Body
                        for row in table.rows[1:]:
                            f.write("| " + " | ".join([cell.text.replace("\n", " ") for cell in row.cells]) + " |\n")
                        f.write("\n")
            print(f"-> Successfully converted {filename} to MD.")
    except Exception as e:
        print(f"-> Failed to convert {filename}: {e}")

print("All conversions complete!")
